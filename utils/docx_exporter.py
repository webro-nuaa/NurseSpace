"""Case → .docx 导出器，输出格式与 DocxParser 完全对称，导出的文件可直接重新上传。"""

from io import BytesIO
from docx import Document


def export_case_to_docx(case, category, stations, knowledge_stations) -> BytesIO:
    """将案例数据导出为标准 .docx 文件。

    Args:
        case: Case ORM 对象
        category: CaseCategory ORM 对象
        stations: list[Station] — assessment 类型站点（按 order_index 排序）
        knowledge_stations: list[Station] — knowledge 类型站点

    Returns:
        BytesIO — 生成的 .docx 文件流
    """
    doc = Document()

    # ---- 案例指引 ----
    if case.case_guide:
        _add_marker(doc, '【案例指引】')
        for line in case.case_guide.split('\n'):
            _add_body(doc, line.strip())
        _add_marker(doc, '【案例指引结尾】')
        doc.add_paragraph('')

    # ---- 站点（assessment） ----
    for station in stations:
        _add_marker(doc, '【站点】')
        _add_body(doc, station.name or '')
        _add_marker(doc, '【站点结尾】')

        if station.assessment_task:
            _add_marker(doc, '【考核任务】')
            for line in station.assessment_task.split('\n'):
                _add_body(doc, line.strip())
            _add_marker(doc, '【考核任务结尾】')

        if getattr(station, 'condition_report', None):
            _add_marker(doc, '【病情汇报】')
            for line in station.condition_report.split('\n'):
                _add_body(doc, line.strip())
            _add_marker(doc, '【病情汇报结尾】')

        _add_marker(doc, '【问题】')
        for line in (station.question or '').split('\n'):
            _add_body(doc, line.strip())
        _add_marker(doc, '【问题结尾】')

        answers = sorted(station.standard_answers.all(), key=lambda a: a.order_index)
        if answers:
            _add_marker(doc, '【回答】')
            for ans in answers:
                _add_marker(doc, '【项】')
                _add_body(doc, ans.answer_item)
                _add_marker(doc, '【项结尾】')
            _add_marker(doc, '【回答结尾】')

        doc.add_paragraph('')

    # ---- 知识拓展 ----
    if knowledge_stations:
        _add_marker(doc, '【知识拓展】')
        for ks in knowledge_stations:
            _add_marker(doc, '【问题】')
            for line in (ks.question or '').split('\n'):
                _add_body(doc, line.strip())
            _add_marker(doc, '【问题结尾】')

            k_answers = sorted(ks.standard_answers.all(), key=lambda a: a.order_index)
            if k_answers:
                _add_marker(doc, '【回答】')
                for ans in k_answers:
                    _add_marker(doc, '【项】')
                    _add_body(doc, ans.answer_item)
                    _add_marker(doc, '【项结尾】')
                _add_marker(doc, '【回答结尾】')
        _add_marker(doc, '【知识拓展结尾】')

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _add_marker(doc, text):
    """添加标记行（如【站点】），清除首行缩进保证对齐。"""
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = None


def _add_body(doc, text):
    """添加正文行。"""
    doc.add_paragraph(text)
