"""Tests for DOCX parser: content structure extraction, file parsing, batch processing."""
import os
import tempfile
import pytest
from docx import Document


def _create_test_docx(content_paragraphs):
    """Helper: create a temporary .docx file with given paragraph texts."""
    doc = Document()
    for text in content_paragraphs:
        doc.add_paragraph(text)
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


class TestDocxContentParsing:
    def test_parse_case_guide(self, app):
        from utils.docx_parser import DocxParser
        parser = DocxParser()
        result = parser._parse_document_content(_make_doc([
            '【案例指引】',
            '这是一段案例指引内容。',
            '包含多行文本。',
            '【案例指引结尾】',
        ]))
        assert '这是一段案例指引内容。' in result['case_guide']
        assert '包含多行文本。' in result['case_guide']

    def test_parse_station_basic(self, app):
        from utils.docx_parser import DocxParser
        parser = DocxParser()
        result = parser._parse_document_content(_make_doc([
            '【站点】',
            '生命体征评估',
            '【考核任务】',
            '评估患者的生命体征。',
            '【考核任务结尾】',
            '【问题】',
            '如何正确测量血压？',
            '【问题结尾】',
            '【回答】',
            '【项】',
            '选择合适的袖带尺寸。',
            '【项结尾】',
            '【项】',
            '袖带下缘距肘窝2-3cm。',
            '【项结尾】',
            '【回答结尾】',
            '【站点结尾】',
        ]))
        assert len(result['stations']) == 1
        s = result['stations'][0]
        assert s['name'] == '生命体征评估'
        assert '生命体征' in s['assessment_task']
        assert '血压' in s['question']
        assert len(s['answers']) == 2
        assert '袖带尺寸' in s['answers'][0]

    def test_parse_multiple_stations(self, app):
        from utils.docx_parser import DocxParser
        parser = DocxParser()
        result = parser._parse_document_content(_make_doc([
            '【站点】', '站点A',
            '【考核任务】', '任务A', '【考核任务结尾】',
            '【问题】', '问题A？', '【问题结尾】',
            '【站点结尾】',
            '【站点】', '站点B',
            '【考核任务】', '任务B', '【考核任务结尾】',
            '【问题】', '问题B？', '【问题结尾】',
            '【站点结尾】',
        ]))
        assert len(result['stations']) == 2
        assert result['stations'][0]['name'] == '站点A'
        assert result['stations'][1]['name'] == '站点B'

    def test_parse_extended_knowledge(self, app):
        from utils.docx_parser import DocxParser
        parser = DocxParser()
        result = parser._parse_document_content(_make_doc([
            '【站点】', '站点X',
            '【考核任务】', '任务X', '【考核任务结尾】',
            '【问题】', '问题X？', '【问题结尾】',
            '【站点结尾】',
            '【知识拓展】',
            '【问题】',
            '拓展问题1？',
            '【问题结尾】',
            '【回答】',
            '拓展回答1。',
            '【回答结尾】',
            '【知识拓展结尾】',
        ]))
        assert len(result['extended_knowledge']) >= 1
        ek = result['extended_knowledge'][0]
        assert '拓展问题1' in ek['question']
        assert '拓展回答1' in ek['answer']

    def test_parse_knowledge_with_items(self, app):
        from utils.docx_parser import DocxParser
        parser = DocxParser()
        result = parser._parse_document_content(_make_doc([
            '【站点】', '站点Y',
            '【考核任务】', '任务Y', '【考核任务结尾】',
            '【问题】', '问题Y？', '【问题结尾】',
            '【站点结尾】',
            '【知识拓展】',
            '【问题】', '拓展问题？', '【问题结尾】',
            '【回答】',
            '【项】', '要点1', '【项结尾】',
            '【项】', '要点2', '【项结尾】',
            '【回答结尾】',
            '【知识拓展结尾】',
        ]))
        assert len(result['extended_knowledge']) >= 1
        ek = result['extended_knowledge'][0]
        assert ek.get('items') is not None

    def test_empty_document(self, app):
        from utils.docx_parser import DocxParser
        parser = DocxParser()
        result = parser._parse_document_content(_make_doc([]))
        assert result['case_guide'] == ''
        assert result['stations'] == []

    def test_no_marks(self, app):
        """Paragraphs without 【】 marks should be ignored."""
        from utils.docx_parser import DocxParser
        parser = DocxParser()
        result = parser._parse_document_content(_make_doc([
            '普通文本，没有标记。',
            '另一段普通文本。',
        ]))
        assert result['case_guide'] == ''
        assert result['stations'] == []


class TestFileParsing:
    def test_parse_file_creates_case(self, app):
        from utils.docx_parser import DocxParser
        from models import Case, db
        parser = DocxParser()
        path = _create_test_docx([
            '【案例指引】', '测试指引内容。', '【案例指引结尾】',
            '【站点】', '测试站点',
            '【考核任务】', '测试任务描述。', '【考核任务结尾】',
            '【问题】', '测试问题？', '【问题结尾】',
            '【回答】',
            '【项】', '答案要点1', '【项结尾】',
            '【回答结尾】',
            '【站点结尾】',
        ])
        # Rename to match expected pattern 【类别】标题.docx
        target = path.replace('.docx', '') + '__tmp.docx'
        os.rename(path, target)
        dirn = os.path.dirname(target)
        fname = '【儿科】测试解析案例.docx'
        final = os.path.join(dirn, fname)
        os.rename(target, final)

        try:
            with app.app_context():
                case = parser.parse_file(final)
                assert case is not None
                assert case.title == '测试解析案例'
                assert case.category.name == '儿科'
                assert case.case_guide is not None
                assert case.stations.count() >= 1
        finally:
            os.unlink(final)

    def test_parse_file_detects_duplicate(self, app):
        from utils.docx_parser import DocxParser
        from models import Case, db
        parser = DocxParser()
        path = _create_test_docx([
            '【案例指引】', '重复案例指引。', '【案例指引结尾】',
            '【站点】', '重复站点',
            '【考核任务】', '任务。', '【考核任务结尾】',
            '【问题】', '问题？', '【问题结尾】',
            '【站点结尾】',
        ])
        target = path.replace('.docx', '') + '__dup.docx'
        os.rename(path, target)
        dirn = os.path.dirname(target)
        fname = '【内科】重复测试.docx'
        final = os.path.join(dirn, fname)
        os.rename(target, final)

        try:
            with app.app_context():
                c1 = parser.parse_file(final)
                c2 = parser.parse_file(final)
                assert c1.id == c2.id  # Same case returned
        finally:
            os.unlink(final)

    def test_parse_file_bad_filename(self, app):
        from utils.docx_parser import DocxParser
        parser = DocxParser()
        path = _create_test_docx(['【案例指引】', '内容', '【案例指引结尾】'])
        # Filename has no 【】 category marker
        dirn = os.path.dirname(path)
        fname = 'no_category_marker.docx'
        final = os.path.join(dirn, fname)
        os.rename(path, final)
        try:
            with app.app_context():
                with pytest.raises(Exception) as exc:
                    parser.parse_file(final)
                assert '类别' in str(exc.value) or '无法从文件名' in str(exc.value)
        finally:
            os.unlink(final)


class TestBatchParsing:
    def test_batch_parse_directory(self, app):
        from utils.docx_parser import DocxParser
        import tempfile
        parser = DocxParser()
        d = tempfile.mkdtemp()
        try:
            # Create 2 valid .docx files
            for i in range(2):
                doc = Document()
                doc.add_paragraph('【案例指引】')
                doc.add_paragraph(f'批量案例{i}指引。')
                doc.add_paragraph('【案例指引结尾】')
                doc.add_paragraph('【站点】')
                doc.add_paragraph(f'站点{i}')
                doc.add_paragraph('【考核任务】')
                doc.add_paragraph('任务描述。')
                doc.add_paragraph('【考核任务结尾】')
                doc.add_paragraph('【问题】')
                doc.add_paragraph(f'问题{i}？')
                doc.add_paragraph('【问题结尾】')
                doc.add_paragraph('【站点结尾】')
                fname = f'【外科】批量案例{i}.docx'
                doc.save(os.path.join(d, fname))

            with app.app_context():
                result = parser.batch_parse_directory(d)
                assert result['success_count'] == 2
                assert result['error_count'] == 0
                assert len(result['results']) == 2
        finally:
            for f in os.listdir(d):
                os.unlink(os.path.join(d, f))
            os.rmdir(d)

    def test_batch_skip_temp_files(self, app):
        from utils.docx_parser import DocxParser
        import tempfile
        parser = DocxParser()
        d = tempfile.mkdtemp()
        try:
            # Create a ~temp file that should be skipped
            doc = Document()
            doc.add_paragraph('【案例指引】')
            doc.add_paragraph('指引。')
            doc.add_paragraph('【案例指引结尾】')
            doc.save(os.path.join(d, '~temp_skip.docx'))

            with app.app_context():
                result = parser.batch_parse_directory(d)
                assert result['success_count'] == 0
        finally:
            for f in os.listdir(d):
                os.unlink(os.path.join(d, f))
            os.rmdir(d)


# ---- helpers ----

def _make_doc(paragraphs):
    """Create an in-memory python-docx Document from a list of paragraph texts."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    return doc
